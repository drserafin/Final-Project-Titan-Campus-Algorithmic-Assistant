"""
ui/notes_search.py — Notes Search Engine module UI.
 
Users load PDF / DOCX / TXT files, enter a search pattern,
choose an algorithm (Naive, Rabin-Karp, KMP, or ALL), and
view match indices + timing comparisons.
"""
 
import tkinter as tk
from tkinter import filedialog, messagebox
from pathlib import Path
 
from ui.theme import (COLORS, FONT_BODY, FONT_SUBHEAD, FONT_SMALL,
                      section_header, card, accent_button, ghost_button,
                      output_box, write_output)
 
from algorithms.string_match import naive_search, rabin_karp, kmp_search
 
# Pre-loaded test assets
ASSETS_DIR = Path(__file__).parent.parent / "assets"
TEST_FILES  = {
    "test_notes.pdf":  ASSETS_DIR / "test_notes.pdf",
    "syllabus.docx":   ASSETS_DIR / "syllabus.docx",
}
 
# How many surrounding characters to show for each match context snippet
CONTEXT_RADIUS = 40
 
 
class NotesSearchFrame(tk.Frame):
    """Module 3 — Naive · Rabin-Karp · KMP"""
 
    def __init__(self, parent):
        super().__init__(parent, bg=COLORS["bg_dark"])
        self._loaded_text: str = ""
        self._loaded_file: str = ""
        self._build()
 
    # ── Layout ──────────────────────────────────────────────────────────────
 
    def _build(self):
        section_header(self,
                        "🔍  Notes Search Engine",
                        "Pattern matching across PDF · DOCX · TXT files")
 
        body = tk.Frame(self, bg=COLORS["bg_dark"])
        body.pack(fill="both", expand=True, padx=24, pady=(0, 20))
        body.columnconfigure(0, weight=0, minsize=280)
        body.columnconfigure(1, weight=1)
        body.rowconfigure(0, weight=1)
 
        self._build_controls(body)
        self._build_output_panel(body)
 
    def _build_controls(self, parent):
        pnl = card(parent)
        pnl.grid(row=0, column=0, sticky="nsew", padx=(0, 12))
 
        # ── File loading section ─────────────────────────────────────────
        tk.Label(pnl, text="Load File", font=FONT_SUBHEAD,
                 fg=COLORS["accent"], bg=COLORS["bg_panel"],
                 anchor="w", padx=12).pack(fill="x", pady=10)
 
        tk.Frame(pnl, bg=COLORS["border"], height=1).pack(fill="x")
 
        accent_button(pnl, "📂  Browse File", command=self._browse_file).pack(
            fill="x", padx=12, pady=(10, 4))
 
        # Quick-load test assets
        tk.Label(pnl, text="Quick Load:", font=FONT_SMALL,
                 fg=COLORS["text_secondary"], bg=COLORS["bg_panel"],
                 anchor="w", padx=12).pack(fill="x")
 
        for name, path in TEST_FILES.items():
            ghost_button(pnl, f"  {name}",
                         command=lambda p=path: self._load_file(p)).pack(
                fill="x", padx=12, pady=2)
 
        # Loaded file indicator
        self._file_label = tk.Label(
            pnl, text="No file loaded",
            font=FONT_SMALL, fg=COLORS["text_secondary"],
            bg=COLORS["bg_panel"], anchor="w",
            padx=12, wraplength=240)
        self._file_label.pack(fill="x", pady=6)
 
        tk.Frame(pnl, bg=COLORS["border"], height=1).pack(fill="x", pady=8)
 
        # ── Pattern input ─────────────────────────────────────────────────
        tk.Label(pnl, text="Search Pattern", font=FONT_SUBHEAD,
                 fg=COLORS["accent"], bg=COLORS["bg_panel"],
                 anchor="w", padx=12).pack(fill="x", pady=(4, 2))
 
        self._pattern_entry = tk.Entry(
            pnl, font=FONT_BODY, width=22,
            bg=COLORS["bg_hover"], fg=COLORS["text_primary"],
            insertbackground=COLORS["accent"],
            relief="flat", bd=4)
        self._pattern_entry.pack(fill="x", padx=12, pady=(0, 6))
        self._pattern_entry.bind("<Return>", lambda e: self._run_all())
 
        # Case-sensitive toggle
        self._case_var = tk.BooleanVar(value=False)
        tk.Checkbutton(
            pnl,
            text="Case-sensitive",
            variable=self._case_var,
            font=FONT_SMALL,
            fg=COLORS["text_secondary"],
            bg=COLORS["bg_panel"],
            activebackground=COLORS["bg_panel"],
            activeforeground=COLORS["text_primary"],
            selectcolor=COLORS["bg_hover"],
            anchor="w",
        ).pack(fill="x", padx=12, pady=(0, 10))
 
        tk.Frame(pnl, bg=COLORS["border"], height=1).pack(fill="x")
 
        # ── Algorithm selection ───────────────────────────────────────────
        tk.Label(pnl, text="Algorithm", font=FONT_SUBHEAD,
                 fg=COLORS["accent"], bg=COLORS["bg_panel"],
                 anchor="w", padx=12).pack(fill="x", pady=(10, 4))
 
        algos = [
            ("Naive Search",    self._run_naive),
            ("Rabin–Karp",      self._run_rk),
            ("KMP",             self._run_kmp),
            ("ALL  (Compare)",  self._run_all),
        ]
        for label, cmd in algos:
            accent_button(pnl, label, command=cmd).pack(
                fill="x", padx=12, pady=3)
 
        tk.Frame(pnl, bg=COLORS["border"], height=1).pack(fill="x", pady=8)
        ghost_button(pnl, "✕  Clear", command=self._clear).pack(
            fill="x", padx=12, pady=(0, 12))
 
    def _build_output_panel(self, parent):
        right = tk.Frame(parent, bg=COLORS["bg_dark"])
        right.grid(row=0, column=1, sticky="nsew")
        right.rowconfigure(0, weight=1)
        right.rowconfigure(1, weight=0, minsize=160)
        right.columnconfigure(0, weight=1)
 
        # Main output
        out_card = card(right)
        out_card.grid(row=0, column=0, sticky="nsew", pady=(0, 10))
 
        tk.Label(out_card, text="Search Results", font=FONT_SUBHEAD,
                 fg=COLORS["accent"], bg=COLORS["bg_panel"],
                 anchor="w", padx=12).pack(fill="x", pady=6)
 
        tk.Frame(out_card, bg=COLORS["border"], height=1).pack(fill="x")
        self._result_output = output_box(out_card, height=16)
 
        # Timing comparison panel
        timing_card = card(right)
        timing_card.grid(row=1, column=0, sticky="nsew")
 
        tk.Label(timing_card, text="Timing Comparison (μs)", font=FONT_SUBHEAD,
                 fg=COLORS["accent"], bg=COLORS["bg_panel"],
                 anchor="w", padx=12).pack(fill="x", pady=6)
 
        tk.Frame(timing_card, bg=COLORS["border"], height=1).pack(fill="x")
        self._timing_output = output_box(timing_card, height=5)
 
    # ── File loading ─────────────────────────────────────────────────────────
 
    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Select a file",
            filetypes=[("All supported", "*.txt *.pdf *.docx"),
                       ("Text", "*.txt"),
                       ("PDF", "*.pdf"),
                       ("Word", "*.docx")])
        if path:
            self._load_file(Path(path))
 
    def _load_file(self, path: Path):
        if not path.exists():
            messagebox.showerror("File Not Found",
                                 f"{path.name} not found in assets/")
            return
        suffix = path.suffix.lower()
        try:
            if suffix == ".txt":
                self._loaded_text = path.read_text(encoding="utf-8", errors="ignore")
            elif suffix == ".pdf":
                self._loaded_text = self._read_pdf(path)
            elif suffix == ".docx":
                self._loaded_text = self._read_docx(path)
            else:
                messagebox.showerror("Unsupported", f"File type '{suffix}' not supported.")
                return
 
            self._loaded_file = path.name
            self._file_label.configure(
                text=f"✔ {path.name}  ({len(self._loaded_text):,} chars)",
                fg=COLORS["success"])
            write_output(self._result_output,
                f"Loaded: {path.name}\n"
                f"Characters: {len(self._loaded_text):,}\n"
                f"Words (approx): {len(self._loaded_text.split()):,}\n\n"
                "Ready to search.")
        except Exception as exc:
            messagebox.showerror("Load Error", str(exc))
 
    def _read_pdf(self, path: Path) -> str:
        try:
            import PyPDF2
            text = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
            return "\n".join(text)
        except ImportError:
            return "[PyPDF2 not installed — run: pip install PyPDF2]"
 
    def _read_docx(self, path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(path))
            lines = []
 
            # Top-level paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
 
            # Table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text)
 
            return "\n".join(lines)
        except ImportError:
            return "[python-docx not installed — run: pip install python-docx]"
 
    # ── Helpers ──────────────────────────────────────────────────────────────
 
    def _guard(self) -> bool:
        if not self._loaded_text:
            messagebox.showwarning("No File", "Please load a file first.")
            return False
        if not self._pattern_entry.get().strip():
            messagebox.showwarning("No Pattern", "Please enter a search pattern.")
            return False
        return True
 
    def _get_search_text(self) -> str:
        """Return text (and pattern) normalised for case if needed."""
        if self._case_var.get():
            return self._loaded_text
        return self._loaded_text.lower()
 
    def _get_pattern(self) -> str:
        pattern = self._pattern_entry.get().strip()
        if not self._case_var.get():
            return pattern.lower()
        return pattern
 
    def _build_result_text(self, algo_name: str, pattern: str,
                           matches: list[int], elapsed: float) -> str:
        """Format result text with match indices and context snippets."""
        original_pattern = self._pattern_entry.get().strip()
        lines = [
            f"Algorithm : {algo_name}",
            f"Pattern   : '{original_pattern}'",
            f"File      : {self._loaded_file}",
            f"Matches   : {len(matches)}",
            f"Time      : {elapsed:.2f} μs",
            "─" * 50,
        ]
 
        if not matches:
            lines.append("No matches found.")
        else:
            text = self._loaded_text  # use original for display
            lines.append(f"Match indices (first 50 shown):\n")
 
            display_limit = min(50, len(matches))
            for rank, idx in enumerate(matches[:display_limit], 1):
                # Context snippet around match
                ctx_start = max(0, idx - CONTEXT_RADIUS)
                ctx_end   = min(len(text), idx + len(original_pattern) + CONTEXT_RADIUS)
                snippet   = text[ctx_start:ctx_end].replace("\n", " ").replace("\r", "")
 
                # Mark where in the snippet the match starts
                match_offset = idx - ctx_start
                before  = snippet[:match_offset]
                matched = snippet[match_offset: match_offset + len(original_pattern)]
                after   = snippet[match_offset + len(original_pattern):]
 
                prefix = "..." if ctx_start > 0 else ""
                suffix = "..." if ctx_end < len(text) else ""
 
                lines.append(
                    f"[{rank:>3}] index {idx}\n"
                    f"      {prefix}{before}[{matched}]{after}{suffix}\n"
                )
 
            if len(matches) > display_limit:
                lines.append(f"  ... and {len(matches) - display_limit} more matches.")
 
        return "\n".join(lines)
 
    def _build_timing_row(self, name: str, matches: list[int],
                          elapsed: float, width: int = 14) -> str:
        return f"{name:<{width}} {elapsed:>10.2f} μs   {len(matches):>7} match{'es' if len(matches) != 1 else ' '}"
 
    # ── Algorithm runners ────────────────────────────────────────────────────
 
    def _run_naive(self):
        if not self._guard():
            return
        text    = self._get_search_text()
        pattern = self._get_pattern()
 
        matches, elapsed = naive_search(text, pattern)
 
        write_output(self._result_output,
                     self._build_result_text("Naive Search", pattern, matches, elapsed))
        write_output(self._timing_output,
            f"{'Algorithm':<14} {'Time (μs)':>13}   {'Matches':>7}\n"
            f"{'─' * 44}\n"
            + self._build_timing_row("Naive", matches, elapsed))
 
    def _run_rk(self):
        if not self._guard():
            return
        text    = self._get_search_text()
        pattern = self._get_pattern()
 
        matches, elapsed = rabin_karp(text, pattern)
 
        write_output(self._result_output,
                     self._build_result_text("Rabin–Karp", pattern, matches, elapsed))
        write_output(self._timing_output,
            f"{'Algorithm':<14} {'Time (μs)':>13}   {'Matches':>7}\n"
            f"{'─' * 44}\n"
            + self._build_timing_row("Rabin-Karp", matches, elapsed))
 
    def _run_kmp(self):
        if not self._guard():
            return
        text    = self._get_search_text()
        pattern = self._get_pattern()
 
        matches, elapsed = kmp_search(text, pattern)
 
        write_output(self._result_output,
                     self._build_result_text("KMP", pattern, matches, elapsed))
        write_output(self._timing_output,
            f"{'Algorithm':<14} {'Time (μs)':>13}   {'Matches':>7}\n"
            f"{'─' * 44}\n"
            + self._build_timing_row("KMP", matches, elapsed))
 
    def _run_all(self):
        if not self._guard():
            return
        text    = self._get_search_text()
        pattern = self._get_pattern()
 
        naive_matches, naive_time   = naive_search(text, pattern)
        rk_matches,    rk_time      = rabin_karp(text, pattern)
        kmp_matches,   kmp_time     = kmp_search(text, pattern)
 
        original_pattern = self._pattern_entry.get().strip()
 
        # Determine fastest algorithm
        times = {
            "Naive":      naive_time,
            "Rabin-Karp": rk_time,
            "KMP":        kmp_time,
        }
        fastest = min(times, key=times.get)
 
        # Main results panel — show all three result blocks
        result_lines = [
            f"ALL ALGORITHMS  —  pattern: '{original_pattern}'",
            f"File: {self._loaded_file}",
            "═" * 50,
            "",
        ]
 
        for name, matches, elapsed in [
            ("Naive Search", naive_matches, naive_time),
            ("Rabin–Karp",   rk_matches,   rk_time),
            ("KMP",          kmp_matches,  kmp_time),
        ]:
            result_lines.append(
                self._build_result_text(name, pattern, matches, elapsed))
            result_lines.append("")
 
        write_output(self._result_output, "\n".join(result_lines))
 
        # Timing panel
        header = (
            f"{'Algorithm':<14} {'Time (μs)':>13}   {'Matches':>7}\n"
            f"{'─' * 44}\n"
        )
        rows = (
            self._build_timing_row("Naive",      naive_matches, naive_time) + "\n" +
            self._build_timing_row("Rabin-Karp", rk_matches,   rk_time)    + "\n" +
            self._build_timing_row("KMP",        kmp_matches,  kmp_time)   + "\n" +
            f"{'─' * 44}\n"
            f"Fastest: {fastest}  ({times[fastest]:.2f} μs)\n"
            f"All algorithms found {len(naive_matches)} match"
            f"{'es' if len(naive_matches) != 1 else ''}."
        )
        write_output(self._timing_output, header + rows)
 
    def _clear(self):
        self._loaded_text = ""
        self._loaded_file = ""
        self._file_label.configure(text="No file loaded",
                                   fg=COLORS["text_secondary"])
        self._pattern_entry.delete(0, "end")
        write_output(self._result_output, "")
        write_output(self._timing_output, "")